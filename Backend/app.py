import streamlit as st
st.set_page_config(page_title="Dr Prepper Chatbot", page_icon="🏥", layout="wide")

# Custom CSS to improve layout
st.markdown("""
<style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .stApp {
        margin-bottom: 0px;
    }
    .stSidebar {
        background-color: #2E2E2E;
    }
    .stSidebar [data-testid="stMarkdownContainer"] {
        color: white;
    }
    .stButton button {
        width: 100%;
    }
    .sidebar-stage {
        padding: 10px;
        border-radius: 5px;
        margin-bottom: 10px;
        cursor: pointer;
        transition: background-color 0.3s;
        width: 100%;
        text-align: left;
        display: block;
    }
    .sidebar-stage.current {
        background-color: orange;
        color: white;
    }
    .sidebar-stage.completed {
        background-color: white;
        color: black;
    }
    .sidebar-stage.upcoming {
        background-color: #f0f0f0;
        color: #888888;
        cursor: not-allowed;
    }
    .stButton.orange button {
        background-color: orange;
        color: white;
    }
    /* Chat input styling */
    .stChatInputContainer {
        padding: 10px;
        background-color: #f0f0f0;
        border-radius: 10px;
        border: 2px solid #4CAF50;
        margin-top: 20px;
    }

    .stChatInputContainer textarea {
        border: none !important;
        background-color: white !important;
        color: #333 !important;
        font-size: 16px !important;
        padding: 10px !important;
        border-radius: 5px !important;
    }

    .stChatInputContainer .stButton > button {
        border-radius: 20px !important;
        background-color: #4CAF50 !important;
        color: white !important;
    }
</style>
""", unsafe_allow_html=True)

# Force Streamlit to use port 8502
import os
os.environ['STREAMLIT_SERVER_PORT'] = '8502'

import time
import json
from openai import OpenAI
import io
from dotenv import load_dotenv
load_dotenv()

if not os.getenv("OPENAI_API_KEY"):
    st.error("OpenAI API key is not set. Please set the OPENAI_API_KEY environment variable.")
    st.stop()

try:
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
except Exception as e:
    st.error(f"Error initializing OpenAI client: {str(e)}")
    st.stop()

ASSISTANT_IDS = {
    "stage1": "asst_a9s6wbqUHXkbIX2vTj5DstO1",
    "stage2": "asst_WfpGAAkD9g0CHzPeHk3FZUD5",
    "stage3": "asst_JeCRuyTbKUi0P3gQh1mTT4yU",
    "stage4": "asst_VI2qxTfRSjFh7SHnEQdH20Lu",
    "stage5": "asst_VkefuaYBuipFKCtxf7cx3fsj"
}

STAGE_TITLES = {
    "stage0": "Fill this out to begin",
    "stage1": "Stage 1: Initial Assessment",
    "stage2": "Stage 2: Possible Diagnoses",
    "stage3": "Stage 3: Probability of Diagnoses",
    "stage4": "Stage 4: Treatment Options",
    "stage5": "Stage 5: Summary for your doctor"
}

# Import SerpAPI library
from serpapi import GoogleSearch

# Get SerpAPI key from environment variable
SERPAPI_API_KEY = os.getenv("SERPAPI_API_KEY")
if not SERPAPI_API_KEY:
    st.error("SerpAPI key is not set. Please set the SERPAPI_API_KEY environment variable.")
    st.stop()

# Define SerpAPI search function
def search_google(query, search_type=None):
    try:
        params = {
            "engine": "google",
            "q": query,
            "api_key": SERPAPI_API_KEY,
            "num": 5  # Limit to 5 results
        }
        search = GoogleSearch(params)
        results = search.get_dict()
        st.write(f"Search results for '{query}' (type: {search_type}):", results)
        
        # Extract and return relevant information
        organic_results = results.get('organic_results', [])
        extracted_results = [
            {
                'title': result.get('title'),
                'link': result.get('link'),
                'snippet': result.get('snippet')
            }
            for result in organic_results
        ]
        return extracted_results
    except Exception as e:
        st.error(f"Error in Google search: {str(e)}")
        return {"error": str(e)}

def generate_response(thread_id, assistant_id, prompt, stage):
    try:
        stage_instructions = {
            "stage1": "Provide an initial assessment based on the user's information.",
            "stage2": "Determine the top 5 possible diagnoses based on the information provided.",
            "stage3": "Analyze the 5 possible diagnoses and rank them from most to least likely. Use your knowledge and the context provided to estimate probabilities. Present this in a markdown table format. Do not use any external data sources.",
            "stage4": "Provide the top 3 treatment options for each diagnosis based on general medical knowledge.",
            "stage5": "Summarize all information for the doctor."
        }
        
        # For stage 3, include a summary of previous diagnoses only
        if stage == "stage3":
            previous_messages = client.beta.threads.messages.list(thread_id=thread_id)
            diagnoses = next((msg.content[0].text.value for msg in reversed(previous_messages.data) if msg.role == "assistant" and "Possible Diagnoses:" in msg.content[0].text.value), "")
            context = f"Previous diagnoses:\n{diagnoses}\n\n"
            full_prompt = f"{context}Current stage: {stage}. Instructions: {stage_instructions.get(stage, '')}. User input: {prompt}"
        else:
            full_prompt = f"Current stage: {stage}. Instructions: {stage_instructions.get(stage, '')}. User input: {prompt}"
        
        client.beta.threads.messages.create(
            thread_id=thread_id,
            role="user",
            content=full_prompt
        )
        
        run = client.beta.threads.runs.create(
            thread_id=thread_id,
            assistant_id=assistant_id
        )
        
        # Add a timeout mechanism
        start_time = time.time()
        timeout = 120  # 120 seconds timeout
        
        while run.status not in ["completed", "failed"]:
            if time.time() - start_time > timeout:
                st.error("Response generation timed out. Please try again.")
                return None
            time.sleep(1)  # Increase sleep time to reduce API calls
            run = client.beta.threads.runs.retrieve(thread_id=thread_id, run_id=run.id)
            
        if run.status == "failed":
            st.error(f"Run failed: {run.last_error}")
            return None

        messages = client.beta.threads.messages.list(thread_id=thread_id)
        for message in messages:
            if message.role == "assistant":
                return {"role": "assistant", "content": message.content[0].text.value, "stage": stage}

    except Exception as e:
        st.error(f"An error occurred while generating the response: {str(e)}")
        return None

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import markdown2
from bs4 import BeautifulSoup

def add_paragraph_with_style(doc, text, style):
    paragraph = doc.add_paragraph(text)
    paragraph.style = style

def process_list_items(doc, items, style):
    for item in items:
        add_paragraph_with_style(doc, item.get_text(), style)
        nested_ul = item.find('ul')
        nested_ol = item.find('ol')
        if nested_ul:
            process_list_items(doc, nested_ul.find_all('li'), 'List Bullet')
        if nested_ol:
            process_list_items(doc, nested_ol.find_all('li'), 'List Number')

def create_word_doc(summary):
    doc = Document()
    
    # Add a title
    doc.add_heading('Health Summary', 0)

    # Convert markdown to HTML
    html = markdown2.markdown(summary)
    
    # Parse HTML and convert to Word elements
    soup = BeautifulSoup(html, 'html.parser')
    
    for element in soup:
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            doc.add_heading(element.get_text(), level)
        elif element.name == 'p':
            doc.add_paragraph(element.get_text())
        elif element.name == 'ul':
            process_list_items(doc, element.find_all('li'), 'List Bullet')
        elif element.name == 'ol':
            process_list_items(doc, element.find_all('li'), 'List Number')
        elif element.name == 'table':
            data = []
            for tr in element.find_all('tr'):
                row = []
                for td in tr.find_all(['td', 'th']):
                    row.append(td.get_text())
                data.append(row)
            table = doc.add_table(rows=len(data), cols=len(data[0]))
            table.style = 'Table Grid'
            for i, row in enumerate(data):
                for j, cell in enumerate(row):
                    table.cell(i, j).text = cell

    # Save the document to a BytesIO object
    doc_output = io.BytesIO()
    doc.save(doc_output)
    doc_output.seek(0)
    return doc_output

# Add this near the top of your file, after the imports
if 'disclaimer_accepted' not in st.session_state:
    st.session_state.disclaimer_accepted = False

# Add this function after the imports and before the main code
def show_disclaimer():
    disclaimer_text = """
    **Disclaimer**

    1. This tool uses AI-generated content and is not a substitute for professional medical advice. It is designed to help you prepare for your doctor's visit. Only a qualified healthcare professional can provide accurate medical advice.

    2. While we do not store your data or request personal information, any information you provide is processed by OpenAI. Although OpenAI is considered secure and trusted, there is always a potential risk of data breaches. To minimize risk, please do not provide any personally identifiable information.

    By using this site, you acknowledge that:
    - This is not professional medical advice
    - You understand the potential risks associated with data processing by third parties
    - You will not provide any personally identifiable information (e.g., name, email, phone number, medical records)
    - You use this site at your own risk and discretion
    """
    
    st.markdown(disclaimer_text)
    
    col1, col2 = st.columns(2)
    with col1:
        agree_medical = st.checkbox("I understand this is not medical advice")
    with col2:
        agree_privacy = st.checkbox("I agree not to provide any personal identifiable information")
    
    if st.button("I Agree", disabled=not (agree_medical and agree_privacy)):
        st.session_state.disclaimer_accepted = True
        st.rerun()  # Use st.rerun() instead of st.experimental_rerun()
    
    if not st.session_state.disclaimer_accepted:
        st.stop()

# Add this check right after the above function
if 'disclaimer_accepted' not in st.session_state or not st.session_state.disclaimer_accepted:
    show_disclaimer()

st.title("Dr Prepper Chatbot")

if "stage" not in st.session_state:
    st.session_state.stage = "stage0"

if "thread_id" not in st.session_state:
    st.session_state.thread_id = client.beta.threads.create().id

if "messages" not in st.session_state:
    st.session_state.messages = []

if "user_info" not in st.session_state:
    st.session_state.user_info = None

if "max_stage" not in st.session_state:
    st.session_state.max_stage = 0

st.sidebar.title("Stages")

def sidebar_stage_button(title, stage, current_stage, max_stage):
    stage_num = int(stage[-1]) if stage != "stage0" else 0
    current_num = int(current_stage[-1]) if current_stage != "stage0" else 0
    
    if stage_num == current_num:
        st.sidebar.markdown(f'<div class="sidebar-stage current">{title}</div>', unsafe_allow_html=True)
    elif stage_num < current_num:
        if st.sidebar.button(title, key=f"sidebar_{stage}", help="Click to go to this stage"):
            st.session_state.stage = stage
            st.session_state.max_stage = stage_num
            st.rerun()
    elif stage_num <= max_stage:
        st.sidebar.markdown(f'<div class="sidebar-stage completed">{title}</div>', unsafe_allow_html=True)
    else:
        st.sidebar.markdown(f'<div class="sidebar-stage upcoming">{title}</div>', unsafe_allow_html=True)

# Display stages in the sidebar
for stage, title in STAGE_TITLES.items():
    sidebar_stage_button(title, stage, st.session_state.stage, st.session_state.max_stage)

if st.session_state.stage == "stage0":
    st.write("Please fill out this form to begin:")
    with st.form("user_info_form"):
        health_issue = st.text_input("What is the single health related issue you want to focus on today?")
        issue_duration = st.text_input("How long have you had this issue?")
        resolution_attempts = st.text_area("What have you done to try to resolve this issue? List surgeries, medication, therapy, counseling, etc.")
        family_history = st.text_input("Does this issue run in your family history? If so, what percent of family members have this issue?")
        birth_year = st.number_input("What year were you born?", min_value=1900, max_value=2023, step=1)
        exercise_habits = st.text_input("How many days a week do you exercise and for how long?")
        diet_rating = st.slider("On a scale from 1-10, how healthy do you eat?", 1, 10, 5)
        sleep_hours = st.number_input("How many hours of sleep do you get a night on average?", min_value=0, max_value=24, step=1)
        
        submit_button = st.form_submit_button("Submit")
        if submit_button:
            st.session_state.user_info = {
                "health_issue": health_issue,
                "issue_duration": issue_duration,
                "resolution_attempts": resolution_attempts,
                "family_history": family_history,
                "birth_year": birth_year,
                "exercise_habits": exercise_habits,
                "diet_rating": diet_rating,
                "sleep_hours": sleep_hours
            }
            st.session_state.stage = "stage1"
            st.session_state.max_stage = 1
            initial_prompt = f"User information: {json.dumps(st.session_state.user_info)}. Please provide an initial assessment based on this information and ask for the user's consent to proceed."
            with st.spinner("Processing your information..."):
                response = generate_response(st.session_state.thread_id, ASSISTANT_IDS["stage1"], initial_prompt, "stage1")
            st.session_state.messages.append(response)
            st.rerun()

else:
    # Display chat messages for the current stage
    for message in st.session_state.messages:
        if int(message.get("stage", "0")[-1]) <= int(st.session_state.stage[-1]):
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

    # Continue button
    current_stage = int(st.session_state.stage[-1])
    if current_stage < st.session_state.max_stage:
        next_stage = f"stage{current_stage + 1}"
        next_stage_title = STAGE_TITLES.get(next_stage, "Finish")
        st.markdown('<div class="stButton orange">', unsafe_allow_html=True)
        if st.button(f"Continue to {next_stage_title}", key="continue_button"):
            st.session_state.stage = next_stage
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)
    elif current_stage < 5:
        next_stage = f"stage{current_stage + 1}"
        next_stage_title = STAGE_TITLES.get(next_stage, "Finish")
        st.markdown('<div class="stButton orange">', unsafe_allow_html=True)
        if st.button(f"Continue to {next_stage_title}", key="continue_button"):
            st.session_state.stage = next_stage
            st.session_state.max_stage = current_stage + 1
            summary_prompt = f"Provide a summary for {next_stage_title}. Start your summary with 'Summary for {next_stage_title}:'"
            with st.spinner("Generating summary..."):
                summary = generate_response(st.session_state.thread_id, ASSISTANT_IDS[next_stage], summary_prompt, next_stage)
            if summary:
                st.session_state.messages.append(summary)
                st.rerun()
            else:
                st.error("Failed to generate summary. Please try again.")
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Word document download button for the final stage
    if st.session_state.stage == "stage5":
        if st.button("Download Word Document", key="download_word"):
            # Retrieve the summary content from the assistant's messages
            summary = next((msg['content'] for msg in reversed(st.session_state.messages) if msg['role'] == 'assistant'), None)
            if summary:
                doc = create_word_doc(summary)
                if doc:
                    st.download_button(
                        label="Click here to download Word Document",
                        data=doc,
                        file_name="health_summary.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.error("Word document generation failed. Please try again later.")
            else:
                st.error("No summary available to download.")

# Chat container at the bottom
st.markdown('<div class="chat-container">', unsafe_allow_html=True)

user_input = st.chat_input("Type your message here...")

st.markdown('<div class="button-container">', unsafe_allow_html=True)
col1, col2 = st.columns(2)
with col1:
    pass  # Remove the Reset Chat button

st.markdown('</div>', unsafe_allow_html=True)
st.markdown('</div>', unsafe_allow_html=True)

if user_input:
    st.session_state.messages.append({"role": "user", "content": user_input, "stage": st.session_state.stage})
    with st.spinner("Generating response..."):
        response = generate_response(st.session_state.thread_id, ASSISTANT_IDS[st.session_state.stage], user_input, st.session_state.stage)
    if response:
        st.session_state.messages.append(response)
        st.rerun()

